import os
import re
import io
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build

# For creating Word documents
from docx import Document
from docx.shared import Pt

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading

# If modifying these scopes, delete your existing token.json.
SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']

###############################################################################
# ADDED: Dictionary for converting Western digits (0-9) to Bangla digits (০-৯)
###############################################################################
BENGALI_DIGITS = {
    '0': '০',
    '1': '১',
    '2': '২',
    '3': '৩',
    '4': '৪',
    '5': '৫',
    '6': '৬',
    '7': '৭',
    '8': '৮',
    '9': '৯'
}

def to_bengali_numerals(number_str):
    """Convert a string of Western digits to Bangla digits."""
    return "".join(BENGALI_DIGITS.get(ch, ch) for ch in number_str)


class WordExporterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Google Sheets to Word Exporter")

        # GUI variables
        self.credentials_path = tk.StringVar()
        self.token_path = tk.StringVar(value='token.json')
        self.spreadsheet_id = tk.StringVar()
        self.output_docx = tk.StringVar(value='output.docx')
        self.range_name = tk.StringVar(value='Sheet1!A1:H')  # Modify to match your data range

        # Layout
        padding = {'padx': 8, 'pady': 4}
        row_index = 0

        tk.Label(root, text="Credentials JSON:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.credentials_path, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_credentials).grid(row=row_index, column=2, **padding)
        row_index += 1

        tk.Label(root, text="Token JSON:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.token_path, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_token).grid(row=row_index, column=2, **padding)
        row_index += 1

        tk.Label(root, text="Spreadsheet ID:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.spreadsheet_id, width=50).grid(row=row_index, column=1, columnspan=2, sticky='w', **padding)
        row_index += 1

        tk.Label(root, text="Spreadsheet Range:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.range_name, width=50).grid(row=row_index, column=1, columnspan=2, sticky='w', **padding)
        row_index += 1

        tk.Label(root, text="Output Word File:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.output_docx, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_output).grid(row=row_index, column=2, **padding)
        row_index += 1

        self.generate_button = ttk.Button(root, text="Generate Word File", command=self.start_generation_thread)
        self.generate_button.grid(row=row_index, column=0, columnspan=3, pady=10)
        row_index += 1

        # Status Text
        self.status_text = tk.Text(root, height=10, width=70, state='disabled')
        self.status_text.grid(row=row_index, column=0, columnspan=3, padx=10, pady=10)

    def browse_credentials(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if file_path:
            self.credentials_path.set(file_path)

    def browse_token(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if file_path:
            self.token_path.set(file_path)

    def browse_output(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
        )
        if file_path:
            self.output_docx.set(file_path)

    def start_generation_thread(self):
        thread = threading.Thread(target=self.run_generation)
        thread.start()

    def run_generation(self):
        self.set_widgets_state('disabled')
        try:
            creds_file = self.credentials_path.get()
            token_file = self.token_path.get()
            spreadsheet_id_val = self.spreadsheet_id.get()
            output_docx_file = self.output_docx.get()
            range_val = self.range_name.get()

            if not creds_file or not spreadsheet_id_val:
                raise ValueError("Please specify credentials JSON and Spreadsheet ID.")

            self.log_status("Starting to export data to Word...")
            export_data_to_word(
                creds_file,
                token_file,
                spreadsheet_id_val,
                range_val,
                output_docx_file,
                self.log_status
            )
            self.log_status(f"Successfully exported data to {output_docx_file}")
            messagebox.showinfo("Success", f"Data exported successfully to {output_docx_file}")

        except Exception as e:
            self.log_status(f"Error: {e}")
            messagebox.showerror("Error", str(e))
        finally:
            self.set_widgets_state('normal')

    def log_status(self, message):
        self.status_text.config(state='normal')
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.status_text.config(state='disabled')

    def set_widgets_state(self, state):
        for child in self.root.winfo_children():
            if isinstance(child, ttk.Entry) or isinstance(child, ttk.Button):
                child.config(state=state)


def export_data_to_word(credentials_path, token_path, spreadsheet_id, range_name, output_docx_path, log_func):
    """Reads rows from the given Google Sheet and exports them into a Word file with the desired structure."""
    # Authenticate
    creds = authenticate_google_sheets(credentials_path, token_path, log_func)
    service = build('sheets', 'v4', credentials=creds)

    # Read data from Google Sheets
    sheet = service.spreadsheets()
    log_func(f"Fetching data from spreadsheet ID '{spreadsheet_id}' range '{range_name}'...")
    result = sheet.values().get(spreadsheetId=spreadsheet_id, range=range_name).execute()
    values = result.get('values', [])

    if not values:
        log_func("No data found in the specified range.")
        return

    headers = values[0]
    data_rows = values[1:]

    # Create a new Word Document
    document = Document()

    # For each row in the Google Sheet, write the structure:
    # [Serial Number]. [Question]
    #
    # [Board/Institute]
    #
    # ক) [Option ক]
    # খ) [Option খ]
    # গ) [Option গ]
    # ঘ) [Option ঘ]
    #
    # উত্তরঃ [Answer]
    serial_index = safe_index(headers, "Serial Number")
    question_index = safe_index(headers, "Question")
    board_index = safe_index(headers, "Board/Institute")
    option_k_index = safe_index(headers, "Option ক")
    option_kha_index = safe_index(headers, "Option খ")
    option_ga_index = safe_index(headers, "Option গ")
    option_gha_index = safe_index(headers, "Option ঘ")
    answer_index = safe_index(headers, "Answer")

    # If your column names are slightly different, adjust the above calls accordingly.

    for row_num, row in enumerate(data_rows, start=2):
        # Safely get each cell (if missing, defaults to empty string)
        serial = get_cell_value(row, serial_index)
        question = get_cell_value(row, question_index)
        board = get_cell_value(row, board_index)
        option_k = get_cell_value(row, option_k_index)
        option_kha = get_cell_value(row, option_kha_index)
        option_ga = get_cell_value(row, option_ga_index)
        option_gha = get_cell_value(row, option_gha_index)
        answer = get_cell_value(row, answer_index)

        ###################################################################
        # ADDED: Convert Western digits in Serial to Bangla before writing
        ###################################################################
        if serial:
            serial_bangla = to_bengali_numerals(serial)
        else:
            serial_bangla = ""

        # 1. "[Serial]. [Question]"
        #    If no serial or question, skip the line (like original logic).
        if serial or question:
            document.add_paragraph(f"{serial_bangla}. {question}")

        # 2. "[Board/Institute]"
        if board:
            document.add_paragraph(board)

        # 3. "ক) [Option ক]", "খ) [Option খ]", "গ) [Option গ]", "ঘ) [Option ঘ]"
        if option_k:
            document.add_paragraph(f"ক) {option_k}")
        if option_kha:
            document.add_paragraph(f"খ) {option_kha}")
        if option_ga:
            document.add_paragraph(f"গ) {option_ga}")
        if option_gha:
            document.add_paragraph(f"ঘ) {option_gha}")

        # 4. "উত্তরঃ [Answer]"
        if answer:
            document.add_paragraph(f"উত্তরঃ {answer}")

        # Add a blank line for spacing
        document.add_paragraph("")

    # Finally, save the Word document
    document.save(output_docx_path)
    log_func("Word file generation complete.")


def authenticate_google_sheets(credentials_path, token_path, log_func):
    """Authenticate with Google Sheets (read-only)."""
    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
        log_func(f"Loaded credentials from {token_path}.")
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
            log_func("Refreshed expired credentials.")
        else:
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
            creds = flow.run_local_server(port=0)
            log_func("Obtained new credentials via OAuth flow.")
        with open(token_path, 'w', encoding='utf-8') as token_file:
            token_file.write(creds.to_json())
            log_func(f"Saved credentials to {token_path}.")
    return creds


def safe_index(header_list, column_name):
    """Return the index of `column_name` in `header_list`, or -1 if not found."""
    # This helps avoid IndexError if a column is missing
    try:
        return header_list.index(column_name)
    except ValueError:
        return -1


def get_cell_value(row, idx):
    """Return the cell value if index is valid; else empty string."""
    if idx == -1 or idx >= len(row):
        return ""
    return row[idx]


def main():
    root = tk.Tk()
    app = WordExporterGUI(root)
    root.mainloop()


if __name__ == '__main__':
    main()
