import os
import re
import io
import urllib.parse
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build

# For creating Word documents
from docx import Document
from docx.shared import Pt

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading

# If modifying these scopes, delete your existing token.json.
SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']

def is_checked(value):
    """
    Interpret a cell's value as a checkbox-like boolean.
    Returns True if the cell string is 'true', '1', 'yes', etc.
    """
    if not value:
        return False
    return value.strip().lower() in ["true", "1", "yes", "checked"]

def parse_sheet_id_from_url(url: str) -> str:
    """
    If `url` is a Google Sheet link (e.g. 'https://docs.google.com/spreadsheets/d/XXXXXX/edit'),
    extract the portion after '/d/' and before the next '/' as the spreadsheet ID.
    If not found, return `url` as-is.
    """
    # Example link pattern:
    # https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/edit#gid=0
    # We want the part 'SPREADSHEET_ID'
    pattern = re.compile(r"docs\.google\.com/spreadsheets/d/([^/]+)/?")
    match = pattern.search(url)
    if match:
        return match.group(1)
    else:
        return url  # Fallback: assume user already provided a direct ID

class WordExporterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Google Sheets to Word Exporter")

        # GUI variables
        self.credentials_path = tk.StringVar()
        self.token_path = tk.StringVar(value='token.json')
        self.spreadsheet_id = tk.StringVar()  # You can paste either the raw ID or the entire link
        self.range_name = tk.StringVar(value='Sheet1!A1:Z')  # Adjust your data range as needed

        # Optional single output file name (not necessarily used when creating 3 separate files).
        self.output_docx = tk.StringVar(value='output.docx')

        # A "common file name" prefix/suffix for the 3 generated files
        self.common_file_name = tk.StringVar(value='common')

        # NEW: Where to save the files
        self.save_dir = tk.StringVar(value='')  # Will hold the directory path

        # Layout
        padding = {'padx': 8, 'pady': 4}
        row_index = 0

        tk.Label(root, text="Credentials JSON:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.credentials_path, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_credentials).grid(row=row_index, column=2, **padding)
        row_index += 1

        tk.Label(root, text="Token JSON:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.token_path, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_token).grid(row=row_index, column=2, **padding)
        row_index += 1

        # Where the user can paste the full Google Sheet link or just the raw ID
        tk.Label(root, text="Spreadsheet ID or Link:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.spreadsheet_id, width=50).grid(row=row_index, column=1, columnspan=2, sticky='w', **padding)
        row_index += 1

        tk.Label(root, text="Spreadsheet Range:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.range_name, width=50).grid(row=row_index, column=1, columnspan=2, sticky='w', **padding)
        row_index += 1

        tk.Label(root, text="Common File Name:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.common_file_name, width=50).grid(row=row_index, column=1, columnspan=2, sticky='w', **padding)
        row_index += 1

        tk.Label(root, text="Save To Directory:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.save_dir, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_save_dir).grid(row=row_index, column=2, **padding)
        row_index += 1

        tk.Label(root, text="(Optional) Single Word File Name:").grid(row=row_index, column=0, sticky='e', **padding)
        ttk.Entry(root, textvariable=self.output_docx, width=50).grid(row=row_index, column=1, sticky='w', **padding)
        ttk.Button(root, text="Browse", command=self.browse_output).grid(row=row_index, column=2, **padding)
        row_index += 1

        self.generate_button = ttk.Button(root, text="Generate Word Files", command=self.start_generation_thread)
        self.generate_button.grid(row=row_index, column=0, columnspan=3, pady=10)
        row_index += 1

        # Status Text
        self.status_text = tk.Text(root, height=10, width=70, state='disabled')
        self.status_text.grid(row=row_index, column=0, columnspan=3, padx=10, pady=10)

    def browse_credentials(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if file_path:
            self.credentials_path.set(file_path)

    def browse_token(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON Files", "*.json")])
        if file_path:
            self.token_path.set(file_path)

    def browse_save_dir(self):
        """
        Lets the user pick a directory where the .docx files will be saved.
        """
        directory = filedialog.askdirectory()
        if directory:
            self.save_dir.set(directory)

    def browse_output(self):
        """
        Optional single docx. Not necessarily used in generating the 3 separate files,
        but we keep it for your original logic or any future expansions.
        """
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
        )
        if file_path:
            self.output_docx.set(file_path)

    def start_generation_thread(self):
        thread = threading.Thread(target=self.run_generation)
        thread.start()

    def run_generation(self):
        self.set_widgets_state('disabled')
        try:
            creds_file = self.credentials_path.get()
            token_file = self.token_path.get()
            raw_input_id_or_link = self.spreadsheet_id.get().strip()
            range_val = self.range_name.get()
            common_name = self.common_file_name.get()
            save_directory = self.save_dir.get().strip()

            # 1. Validate required fields
            if not creds_file:
                raise ValueError("Please specify credentials JSON.")
            if not raw_input_id_or_link:
                raise ValueError("Please provide a Spreadsheet ID or the full link.")
            if not save_directory:
                raise ValueError("Please select a directory to save the files.")

            # 2. If the user pasted a full link, parse out the ID
            spreadsheet_id_val = parse_sheet_id_from_url(raw_input_id_or_link)

            self.log_status("Starting to export data to 3 separate Word files (Class, Lecture, Quiz)...")
            export_data_to_three_files(
                credentials_path=creds_file,
                token_path=token_file,
                spreadsheet_id=spreadsheet_id_val,
                range_name=range_val,
                log_func=self.log_status,
                common_file_name=common_name,
                save_dir=save_directory
            )
            self.log_status(
                "Successfully exported data to 3 separate Word files:\n"
                f"1) {common_name}_class_slide.docx\n"
                f"2) {common_name}_lecture_sheet.docx\n"
                f"3) {common_name}_quiz.docx\n"
                f"in directory: {save_directory}"
            )
            messagebox.showinfo(
                "Success",
                f"Data exported successfully to:\n\n"
                f"{save_directory}\n\n"
                f"1) {common_name}_class_slide.docx\n"
                f"2) {common_name}_lecture_sheet.docx\n"
                f"3) {common_name}_quiz.docx"
            )

        except Exception as e:
            self.log_status(f"Error: {e}")
            messagebox.showerror("Error", str(e))
        finally:
            self.set_widgets_state('normal')

    def log_status(self, message):
        self.status_text.config(state='normal')
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.status_text.config(state='disabled')

    def set_widgets_state(self, state):
        for child in self.root.winfo_children():
            if isinstance(child, ttk.Entry) or isinstance(child, ttk.Button):
                child.config(state=state)


def export_data_to_three_files(credentials_path, token_path, spreadsheet_id, range_name, log_func, common_file_name, save_dir):
    """
    Reads rows from Google Sheets, then creates 3 docx files in `save_dir`:
      1) {common_file_name}_class_slide.docx
      2) {common_file_name}_lecture_sheet.docx
      3) {common_file_name}_quiz.docx
    
    Based on "For Class Slide", "For Lecture sheet", "For Quiz" columns.
    The structure for each MCQ is:

    [Serial Number]. [Question]

    [Board/Institute]

    ক) [Option ক]
    খ) [Option খ]
    গ) [Option গ]
    ঘ) [Option ঘ]

    উত্তরঃ [Answer]
    """
    # 1. Authenticate
    creds = authenticate_google_sheets(credentials_path, token_path, log_func)
    service = build('sheets', 'v4', credentials=creds)

    # 2. Fetch data
    sheet = service.spreadsheets()
    log_func(f"Fetching data from spreadsheet ID '{spreadsheet_id}' range '{range_name}'...")
    result = sheet.values().get(spreadsheetId=spreadsheet_id, range=range_name).execute()
    values = result.get('values', [])

    if not values:
        log_func("No data found in the specified range.")
        return

    headers = values[0]
    data_rows = values[1:]

    # 3. Create 3 Word documents
    class_slide_doc = Document()
    lecture_doc = Document()
    quiz_doc = Document()

    # 4. Identify column indices
    serial_index = safe_index(headers, "Serial Number")
    question_index = safe_index(headers, "Question")
    class_slide_index = safe_index(headers, "For Class Slide")
    lecture_index = safe_index(headers, "For Lecture sheet")
    quiz_index = safe_index(headers, "For Quiz")
    board_index = safe_index(headers, "Board/Institute")
    option_k_index = safe_index(headers, "Option ক")
    option_kha_index = safe_index(headers, "Option খ")
    option_ga_index = safe_index(headers, "Option গ")
    option_gha_index = safe_index(headers, "Option ঘ")
    answer_index = safe_index(headers, "Answer")

    # 5. Loop through each row and decide where it goes
    for row_num, row in enumerate(data_rows, start=2):
        serial = get_cell_value(row, serial_index)
        question = get_cell_value(row, question_index)

        for_class_slide_val = get_cell_value(row, class_slide_index)
        for_lecture_val = get_cell_value(row, lecture_index)
        for_quiz_val = get_cell_value(row, quiz_index)

        board = get_cell_value(row, board_index)
        option_k = get_cell_value(row, option_k_index)
        option_kha = get_cell_value(row, option_kha_index)
        option_ga = get_cell_value(row, option_ga_index)
        option_gha = get_cell_value(row, option_gha_index)
        answer = get_cell_value(row, answer_index)

        # If the row is checked for "Class Slide", add it to class_slide_doc
        if is_checked(for_class_slide_val):
            add_mcq_to_document(
                doc=class_slide_doc,
                serial=serial,
                question=question,
                board=board,
                option_k=option_k,
                option_kha=option_kha,
                option_ga=option_ga,
                option_gha=option_gha,
                answer=answer
            )

        # If the row is checked for "Lecture sheet", add it to lecture_doc
        if is_checked(for_lecture_val):
            add_mcq_to_document(
                doc=lecture_doc,
                serial=serial,
                question=question,
                board=board,
                option_k=option_k,
                option_kha=option_kha,
                option_ga=option_ga,
                option_gha=option_gha,
                answer=answer
            )

        # If the row is checked for "Quiz", add it to quiz_doc
        if is_checked(for_quiz_val):
            add_mcq_to_document(
                doc=quiz_doc,
                serial=serial,
                question=question,
                board=board,
                option_k=option_k,
                option_kha=option_kha,
                option_ga=option_ga,
                option_gha=option_gha,
                answer=answer
            )

    # 6. Save each document in the specified directory with the common name
    class_slide_path = os.path.join(save_dir, f"{common_file_name}_class_slide.docx")
    lecture_path = os.path.join(save_dir, f"{common_file_name}_lecture_sheet.docx")
    quiz_path = os.path.join(save_dir, f"{common_file_name}_quiz.docx")

    class_slide_doc.save(class_slide_path)
    lecture_doc.save(lecture_path)
    quiz_doc.save(quiz_path)

    log_func(f"Successfully created 3 Word files in: {save_dir}")


def add_mcq_to_document(doc, serial, question, board, option_k, option_kha, option_ga, option_gha, answer):
    """
    Writes the MCQ in this structure:
    [Serial Number]. [Question]

    [Board/Institute]

    ক) [Option ক]
    খ) [Option খ]
    গ) [Option গ]
    ঘ) [Option ঘ]

    উত্তরঃ [Answer]
    """
    # 1. "[Serial]. [Question]"
    line = ""
    if serial or question:
        if serial and question:
            line = f"{serial}. {question}"
        elif serial:
            line = f"{serial}."
        else:
            line = f"{question}"
        doc.add_paragraph(line.strip())

    # 2. "[Board/Institute]"
    if board:
        doc.add_paragraph(board)

    # 3. Options
    if option_k:
        doc.add_paragraph(f"ক) {option_k}")
    if option_kha:
        doc.add_paragraph(f"খ) {option_kha}")
    if option_ga:
        doc.add_paragraph(f"গ) {option_ga}")
    if option_gha:
        doc.add_paragraph(f"ঘ) {option_gha}")

    # 4. "উত্তরঃ [Answer]"
    if answer:
        doc.add_paragraph(f"উত্তরঃ {answer}")

    # Add a blank line
    doc.add_paragraph("")


def authenticate_google_sheets(credentials_path, token_path, log_func):
    """Authenticate with Google Sheets (read-only)."""
    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
        log_func(f"Loaded credentials from {token_path}.")
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
            log_func("Refreshed expired credentials.")
        else:
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
            creds = flow.run_local_server(port=0)
            log_func("Obtained new credentials via OAuth flow.")
        with open(token_path, 'w', encoding='utf-8') as token_file:
            token_file.write(creds.to_json())
            log_func(f"Saved credentials to {token_path}.")
    return creds


def safe_index(header_list, column_name):
    """Return the index of `column_name` in `header_list`, or -1 if not found."""
    try:
        return header_list.index(column_name)
    except ValueError:
        return -1


def get_cell_value(row, idx):
    """Return the cell value if index is valid; else empty string."""
    if idx == -1 or idx >= len(row):
        return ""
    return row[idx]


def main():
    root = tk.Tk()
    app = WordExporterGUI(root)
    root.mainloop()


if __name__ == '__main__':
    main()
